#!/usr/bin/env python3
"""Build legacy-parity research artifacts from the same frozen evaluator payload."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")


def safe_name(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|\s]+', "_", value).strip("_") or "case"


def text(value: Any, fallback: str = "尚待補件") -> str:
    return fallback if value in (None, "", []) else str(value)


def add_table(doc: Document, rows: list[list[Any]]) -> None:
    if not rows:
        doc.add_paragraph("尚待補件")
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = text(value, "")
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = text(value, "")


def save_docx(path: Path, title: str, subtitle: str, sections: list[tuple[str, Any]]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(title, 0)
    doc.add_paragraph(subtitle)
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        if isinstance(body, list) and body and isinstance(body[0], list):
            add_table(doc, body)
        elif isinstance(body, list):
            for item in body or ["尚待補件"]:
                doc.add_paragraph(text(item), style="List Bullet")
        else:
            doc.add_paragraph(text(body))
    doc.add_paragraph("內部研究草稿，非投資建議；未有直接來源的欄位均維持尚待補件。")
    doc.save(path)


def build_dd_tracker(path: Path, data: dict[str, Any]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "DD_Request_Tracker"
    headers = ["ID", "類型", "優先級", "項目／問題", "原因", "Owner", "期限", "狀態", "來源定位"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = PatternFill("solid", fgColor="244A68")
    rows: list[list[Any]] = []
    for index, item in enumerate(data.get("missing_items", []), 1):
        rows.append([f"D{index}", "補件", item.get("priority"), item.get("item"), item.get("reason"), item.get("owner"), item.get("due"), item.get("status"), item.get("locator")])
    for index, question in enumerate(data.get("deck", {}).get("management_questions", []), 1):
        rows.append([f"Q{index}", "管理層問答", "P1", question, "驗證投資命題或風險", "Management", "", "待回答", ""])
    if not rows:
        rows.append(["D1", "補件", "P0", "尚未建立補件清單", "資料不足", "", "", "待建立", ""])
    for row in rows:
        ws.append(row)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    widths = [10, 16, 12, 48, 36, 18, 16, 14, 30]
    for index, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + index)].width = width
    wb.save(path)


def build_markdown(data: dict[str, Any]) -> str:
    deck = data.get("deck", {})
    history = sorted(data.get("financial_history", []), key=lambda row: row.get("year", 0))
    lines = [
        f"# {data['meta']['company']} Investment Memo",
        "",
        f"- Case: {data['meta']['case_id']}",
        f"- Mode: {data.get('mode', 'full')}",
        f"- Data as of: {data['meta']['as_of']}",
        "",
        "## 投資命題與決策邊界",
        "",
        text(deck.get("thesis")),
        "",
        f"目前狀態：{text(deck.get('decision_status'))}",
        "",
        "## 歷史財務摘要",
        "",
        "| 年度 | 營收 | 淨利 | 現金 | 負債 | 來源 |",
        "|---:|---:|---:|---:|---:|---|",
    ]
    for row in history:
        lines.append(f"| {text(row.get('year'))} | {text(row.get('revenue'))} | {text(row.get('net_income'))} | {text(row.get('cash'))} | {text(row.get('debt'))} | {text(row.get('source'))} |")
    for heading, items in (("前置條件", deck.get("conditions", [])), ("主要風險", [row[0] for row in deck.get("risks", [])]), ("缺件", [item.get("item") for item in data.get("missing_items", [])])):
        lines.extend(["", f"## {heading}", ""])
        lines.extend(f"- {text(item)}" for item in (items or ["尚待補件"]))
    lines.extend(["", "內部研究草稿，非投資建議；最終決策由 GP 作成。", ""])
    return "\n".join(lines)


def record_rows(records: list[dict[str, Any]], columns: list[str]) -> list[list[Any]]:
    return [columns] + [[record.get(column) for column in columns] for record in records]


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    try:
        data = json.loads(args.input.read_text(encoding="utf-8"))
        args.output_dir.mkdir(parents=True, exist_ok=True)
        base = safe_name(data["meta"]["company"])
        subtitle = f"Case {data['meta']['case_id']}｜資料截至 {data['meta']['as_of']}｜mode={data.get('mode', 'full')}"
        outputs: dict[str, str] = {}

        memo_md = args.output_dir / f"{base}_Investment_Memo.md"
        memo_md.write_text(build_markdown(data), encoding="utf-8")
        outputs["investment_memo_md"] = str(memo_md)
        memo_docx = args.output_dir / f"{base}_Investment_Memo.docx"
        save_docx(memo_docx, f"{data['meta']['company']} Investment Memo", subtitle, [
            ("投資命題", data.get("deck", {}).get("thesis")),
            ("目前狀態", data.get("deck", {}).get("decision_status")),
            ("前置條件", data.get("deck", {}).get("conditions", [])),
            ("主要風險", [row[0] for row in data.get("deck", {}).get("risks", [])]),
        ])
        outputs["investment_memo_docx"] = str(memo_docx)

        dd_path = args.output_dir / f"{base}_DD_Request_Tracker.xlsx"
        build_dd_tracker(dd_path, data)
        outputs["dd_request_tracker"] = str(dd_path)

        qa_path = args.output_dir / f"{base}_Management_QA.docx"
        save_docx(qa_path, f"{data['meta']['company']} Management Q&A", subtitle, [("待訪談問題", data.get("deck", {}).get("management_questions", []))])
        outputs["management_qa"] = str(qa_path)

        legacy = data.get("legacy_outputs", {})
        interview_path = args.output_dir / f"{base}_Interview_Notes.docx"
        save_docx(interview_path, f"{data['meta']['company']} Interview Notes", subtitle, [("訪談紀錄", record_rows(legacy.get("interview_records", []), ["date", "participants", "question", "answer", "source_locator"]))])
        outputs["interview_notes"] = str(interview_path)

        meeting_path = args.output_dir / f"{base}_Meeting_Minutes.docx"
        save_docx(meeting_path, f"{data['meta']['company']} Meeting Minutes", subtitle, [("會議紀錄", record_rows(legacy.get("meeting_records", []), ["date", "participants", "topic", "decision", "action", "owner"]))])
        outputs["meeting_minutes"] = str(meeting_path)

        notes_path = args.output_dir / f"{base}_Financial_Statement_Notes.docx"
        save_docx(notes_path, f"{data['meta']['company']} Financial Statement Notes", subtitle, [("財報附註重點", record_rows(legacy.get("financial_statement_notes", []), ["topic", "finding", "impact", "source_locator"]))])
        outputs["financial_statement_notes"] = str(notes_path)

        report = {"status": "complete", "artifact_count": len(outputs), "outputs": outputs}
        report_path = args.output_dir / f"{base}_legacy_parity_manifest.json"
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        report["manifest"] = str(report_path)
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return 0
    except (OSError, KeyError, ValueError, json.JSONDecodeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
