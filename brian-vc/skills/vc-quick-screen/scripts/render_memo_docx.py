#!/usr/bin/env python3
"""Render a vc-quick-screen Markdown memo to a neutral, compact DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError as exc:
    if exc.name != "docx":
        raise
    raise SystemExit(
        "ERROR: vc-quick-screen 的 DOCX renderer 需要 python-docx。"
        "在 Codex/ChatGPT 請使用內建 Documents runtime；"
        "獨立執行時請先安裝 brian-vc/requirements.txt"
        "（python -m pip install -r brian-vc/requirements.txt）。"
    ) from exc


ACCENT = "1F4E79"
MUTED = "666666"
LIGHT_FILL = "EAF0F6"
GRID = "B7C3D0"
ASCII_FONT = "Arial"
CJK_FONT = "Microsoft JhengHei"
# Named compact quick-screen override: 0.55-inch side margins yield a
# 7.4-inch content measure. Wider tables reduce wrapping without shrinking
# table text below a readable size.
TABLE_WIDTH_DXA = 10656
TABLE_INDENT_DXA = 120
ICON_LABELS = {
    "🔴": "[紅旗]",
    "🟡": "[注意]",
    "🟢": "[亮點]",
    "⚠️": "[待補]",
    "⚠": "[待補]",
    "✅": "[通過]",
}


def set_run_font(
    run,
    *,
    size: float,
    bold: bool | None = None,
    color: str = "000000",
    italic: bool | None = None,
) -> None:
    run.font.name = ASCII_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ASCII_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top: int = 55, start: int = 90, bottom: int = 55, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def choose_widths(rows: list[list[str]], ncols: int) -> list[int]:
    weights = []
    for col in range(ncols):
        longest = max(
            (len(re.sub(r"\s+", "", row[col])) if col < len(row) else 0)
            for row in rows
        )
        weights.append(max(5, min(longest, 42)))
    total = sum(weights)
    widths = [max(650, round(TABLE_WIDTH_DXA * weight / total)) for weight in weights]
    delta = TABLE_WIDTH_DXA - sum(widths)
    widths[-1] += delta
    return widths


def add_inline(paragraph, text: str, *, size: float = 9.0, color: str = "000000") -> None:
    # A semantic heading such as "🟢 亮點" already names the category; avoid
    # rendering it as the redundant "[亮點] 亮點".
    text = re.sub(r"^\s*🟢\s*(?:亮點|值得保留觀察的\s*3?\s*點)\s*$", "亮點", text)
    text = re.sub(r"^\s*🔴\s*(?:疑慮|疑慮／紅旗|紅旗)(?:\s*3\s*點)?\s*$", "疑慮", text)
    for icon, label in ICON_LABELS.items():
        text = text.replace(icon, label)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]), size=size, color=color)
        set_run_font(paragraph.add_run(match.group(1)), size=size, bold=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(9.0)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    title = doc.styles["Title"]
    title.font.name = ASCII_FONT
    title.font.size = Pt(19)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(ACCENT)
    title._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    title.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after in (
        ("Heading 1", 12.0, 6, 2),
        ("Heading 2", 10.2, 4, 1),
    ):
        style = doc.styles[style_name]
        style.font.name = ASCII_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = ASCII_FONT
        style.font.size = Pt(8.8)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_after = Pt(1)

    if "Quick Screen Note" not in doc.styles:
        note = doc.styles.add_style("Quick Screen Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = doc.styles["Quick Screen Note"]
    note.font.name = ASCII_FONT
    note.font.size = Pt(8.4)
    note.font.italic = True
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    note.paragraph_format.left_indent = Inches(0.15)
    note.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    ncols = max(len(row) for row in rows)
    normalized = [row + [""] * (ncols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    if ncols == 4 and normalized[0][0].strip() == "專家":
        # Named override: keep the six expert-role labels on one line while
        # preserving more width for observation and question content.
        widths = [1800, 3400, 4556, 900]
    else:
        widths = choose_widths(normalized, ncols)
    set_table_geometry(table, widths)

    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_tr_pr.append(repeat)

    for row_index, row in enumerate(normalized):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit")
        row_properties.append(cannot_split)
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, LIGHT_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if row_index == 0:
                paragraph.paragraph_format.keep_with_next = True
            add_inline(
                paragraph,
                value,
                size=7.6 if ncols >= 5 else 8.0,
                color=ACCENT if row_index == 0 else "000000",
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.line_spacing = 0.5


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()
        if not re.fullmatch(r"\|[\s:|-]+\|", raw):
            rows.append([cell.strip() for cell in raw.strip("|").split("|")])
        index += 1
    return rows, index


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def new_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA1000000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"{0xB1000000 + abstract_id:08X}")
    abstract.append(tmpl)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    level.append(suffix)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    # OOXML requires all abstractNum elements to precede num elements.
    # Appending the abstract after existing nums makes Word repair the file and
    # can collapse otherwise separate lists into one continuous sequence.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def render(markdown: Path, output: Path) -> None:
    text = markdown.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Named compact-memo override: keep the complete vQS memo within its
    # explicit three-page delivery target without reducing body text below 9 pt.
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = ""
    set_run_font(
        header.add_run("VC QUICK SCREEN  |  初步研究觀察"),
        size=8,
        bold=True,
        color=MUTED,
    )
    header.paragraph_format.space_after = Pt(0)
    add_page_number(section.footer.paragraphs[0])

    index = 0
    in_numbered_list = False
    current_num_id = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("|"):
            in_numbered_list = False
            rows, index = parse_table(lines, index)
            if rows:
                add_table(doc, rows)
            continue
        if stripped.startswith("# "):
            in_numbered_list = False
            paragraph = doc.add_paragraph(style="Title")
            add_inline(paragraph, stripped[2:].strip(), size=19, color=ACCENT)
        elif stripped.startswith("## "):
            in_numbered_list = False
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline(paragraph, stripped[3:].strip(), size=12.5, color=ACCENT)
        elif stripped.startswith("### "):
            in_numbered_list = False
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, stripped[4:].strip(), size=10.5, color=ACCENT)
        elif stripped.startswith(">"):
            in_numbered_list = False
            paragraph = doc.add_paragraph(style="Quick Screen Note")
            add_inline(paragraph, stripped.lstrip("> ").strip(), size=8.4, color=MUTED)
        elif re.match(r"^\d+\.\s+", stripped):
            if not in_numbered_list:
                current_num_id = new_decimal_numbering(doc)
            in_numbered_list = True
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            apply_numbering(paragraph, current_num_id)
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", stripped), size=8.8)
        elif re.match(r"^[-*]\s+", stripped):
            in_numbered_list = False
            paragraph = doc.add_paragraph(style="List Bullet")
            apply_numbering(paragraph, 1)
            add_inline(paragraph, re.sub(r"^[-*]\s+", "", stripped), size=8.8)
        else:
            in_numbered_list = False
            paragraph = doc.add_paragraph()
            add_inline(paragraph, stripped)
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = re.sub(r"^#\s*", "", lines[0]) if lines else "VC Quick Screen"
    doc.core_properties.subject = "VC quick-screen research observation"
    doc.core_properties.author = "VC Quick Screen"
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    render(args.input_md, args.output_docx)
    print(f"WROTE {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
